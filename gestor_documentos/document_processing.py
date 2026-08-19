from pathlib import Path
from zipfile import BadZipFile
from xml.etree import ElementTree

from docx import Document as WordDocument
from docx.document import Document as WordProcessingDocument
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table
from docx.text.paragraph import Paragraph
from django.utils.html import escape


ALLOWED_EXTENSIONS = {".docx", ".txt", ".md"}


class DocumentProcessingError(Exception):
    pass


def get_extension(filename):
    return Path(filename).suffix.lower()


def validate_extension(filename):
    extension = get_extension(filename)
    if extension not in ALLOWED_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise DocumentProcessingError(
            f"El archivo {filename} no es compatible. Solo se permiten: {allowed}."
        )
    return extension


def convert_uploaded_file_to_markdown(uploaded_file):
    extension = validate_extension(uploaded_file.name)

    if extension == ".docx":
        return _convert_docx_to_markdown(uploaded_file)

    return _convert_text_to_markdown(uploaded_file)


def build_markdown_filename(filename):
    extension = get_extension(filename)
    if extension == ".md":
        return filename
    return f"{filename}.md"


def render_stored_file_to_html(file_field, extension):
    extension = (extension or get_extension(file_field.name or "")).lower()

    if extension == ".docx":
        return _render_docx_to_html(file_field)

    text = _read_text_file(file_field)
    return _render_text_to_html(text)


def _convert_text_to_markdown(uploaded_file):
    uploaded_file.seek(0)
    raw = uploaded_file.read()
    uploaded_file.seek(0)

    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise DocumentProcessingError(
        f"No se pudo leer el contenido del archivo {uploaded_file.name}."
    )


def _convert_docx_to_markdown(uploaded_file):
    uploaded_file.seek(0)
    try:
        document = WordDocument(uploaded_file)
    except (BadZipFile, PackageNotFoundError):
        raise DocumentProcessingError(
            f"El archivo {uploaded_file.name} no es un .docx valido o esta dañado."
        ) from None
    uploaded_file.seek(0)

    blocks = []
    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            _append_paragraph_block(blocks, block)
            continue

        if isinstance(block, Table):
            _append_table_block(blocks, block)

    markdown = "\n\n".join(block for block in blocks if block is not None).strip()
    if not markdown:
        raise DocumentProcessingError(
            f"El archivo {uploaded_file.name} no contiene texto convertible a Markdown."
        )
    return markdown + "\n"


def _render_docx_to_html(file_field):
    file_field.open("rb")
    try:
        try:
            document = WordDocument(file_field)
        except (BadZipFile, PackageNotFoundError):
            raise DocumentProcessingError(
                f"El archivo {file_field.name} no es un .docx valido o esta dañado."
            ) from None
    finally:
        file_field.close()

    blocks = []
    active_list_type = None
    for block in _iter_block_items(document):
        if isinstance(block, Paragraph):
            active_list_type = _append_paragraph_html(blocks, block, active_list_type)
            continue

        if active_list_type:
            blocks.append(f"</{active_list_type}>")
            active_list_type = None
        if isinstance(block, Table):
            blocks.append(_table_to_html(block))

    if active_list_type:
        blocks.append(f"</{active_list_type}>")

    rendered = "".join(blocks).strip()
    if not rendered:
        raise DocumentProcessingError(
            f"El archivo {file_field.name} no contiene texto visible para renderizar."
        )
    return rendered


def _extract_heading_level(style_name):
    for token in style_name.split():
        if token.isdigit():
            return max(1, min(int(token), 6))
    return 1


def _get_style_name(paragraph):
    style = getattr(paragraph, "style", None)
    if style is None:
        return ""
    return (getattr(style, "name", "") or "").lower()


def _iter_block_items(document):
    parent = document.element.body
    for child in parent.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def _append_paragraph_block(blocks, paragraph):
    text = paragraph.text.strip()
    if not text:
        if blocks and blocks[-1] != "":
            blocks.append("")
        return

    style_name = _get_style_name(paragraph)
    if style_name.startswith("heading"):
        level = _extract_heading_level(style_name)
        blocks.append(f"{'#' * level} {text}")
        return

    blocks.append(text)


def _append_paragraph_html(blocks, paragraph, active_list_type):
    text_html = _render_paragraph_runs(paragraph)
    style_name = _get_style_name(paragraph)

    if not text_html.strip():
        if active_list_type:
            blocks.append(f"</{active_list_type}>")
        blocks.append('<div class="viewer-spacer"></div>')
        return None

    list_type = _detect_list_type(paragraph, style_name)
    if list_type:
        if active_list_type != list_type:
            if active_list_type:
                blocks.append(f"</{active_list_type}>")
            blocks.append(f"<{list_type}>")
        blocks.append(f"<li>{text_html}</li>")
        return list_type

    if active_list_type:
        blocks.append(f"</{active_list_type}>")
        active_list_type = None

    if style_name.startswith("heading"):
        level = _extract_heading_level(style_name)
        blocks.append(f"<h{level}>{text_html}</h{level}>")
        return None

    blocks.append(f"<p>{text_html}</p>")
    return None


def _append_table_block(blocks, table):
    rows = []
    for row in table.rows:
        cells = [_normalize_table_cell(cell.text) for cell in row.cells]
        if any(cells):
            rows.append(cells)

    if not rows:
        return

    column_count = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (column_count - len(row)) for row in rows]
    header = normalized_rows[0]
    separator = ["---"] * column_count

    table_lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in normalized_rows[1:]:
        table_lines.append("| " + " | ".join(row) + " |")

    if blocks and blocks[-1] != "":
        blocks.append("")
    blocks.append("\n".join(table_lines))
    blocks.append("")


def _table_to_html(table):
    rows = []
    for row in table.rows:
        cells = [_render_table_cell_html(cell) for cell in row.cells]
        if any(cell.strip() for cell in cells):
            rows.append(cells)

    if not rows:
        return ""

    html = ['<div class="viewer-table-wrap"><table class="viewer-table">']
    header = rows[0]
    html.append("<thead><tr>")
    for cell in header:
        html.append(f"<th>{cell}</th>")
    html.append("</tr></thead>")

    if len(rows) > 1:
        html.append("<tbody>")
        for row in rows[1:]:
            html.append("<tr>")
            for cell in row:
                html.append(f"<td>{cell}</td>")
            html.append("</tr>")
        html.append("</tbody>")

    html.append("</table></div>")
    return "".join(html)


def _normalize_table_cell(text):
    return " ".join((text or "").split()).replace("|", "\\|")


def _read_text_file(file_field):
    file_field.open("rb")
    try:
        raw = file_field.read()
    finally:
        file_field.close()

    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise DocumentProcessingError(
        f"No se pudo leer el contenido del archivo {file_field.name}."
    )


def _render_text_to_html(text):
    lines = text.splitlines()
    if not lines:
        return '<p class="viewer-empty-copy">Este documento no tiene contenido para mostrar.</p>'

    html = []
    for line in lines:
        if line.strip():
            html.append(f"<p>{escape(line)}</p>")
        else:
            html.append('<div class="viewer-spacer"></div>')
    return "".join(html)


def _render_paragraph_runs(paragraph):
    parts = []
    if paragraph.runs:
        for run in paragraph.runs:
            text = escape(run.text).replace("\n", "<br>")
            if not text:
                continue
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"
            parts.append(text)
    else:
        parts.append(escape(paragraph.text).replace("\n", "<br>"))
    return "".join(parts).strip()


def _render_table_cell_html(cell):
    fragments = []
    for paragraph in cell.paragraphs:
        text_html = _render_paragraph_runs(paragraph)
        if text_html:
            fragments.append(text_html)
    return "<br>".join(fragments)


def _detect_list_type(paragraph, style_name):
    if "list bullet" in style_name or "bullet" in style_name:
        return "ul"
    if "list number" in style_name or "number" in style_name:
        return "ol"

    xml = paragraph._p.xml
    if "<w:numPr>" in xml:
        root = ElementTree.fromstring(xml)
        namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        level = root.find(".//w:numPr", namespaces)
        if level is not None:
            return "ol"
    return None
